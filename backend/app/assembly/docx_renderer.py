"""Convert assembled markdown proposal to DOCX via python-docx."""

import io
import re

from docx import Document
from docx.shared import Pt, RGBColor


def render_docx(markdown_text: str, title: str = "Proposal") -> bytes:
    """Render markdown proposal to DOCX bytes."""
    doc = Document()
    _setup_styles(doc)
    _parse_markdown_to_docx(doc, markdown_text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color) in enumerate([
        (20, (0x1A, 0x36, 0x5D)),
        (16, (0x2C, 0x52, 0x82)),
        (13, (0x2D, 0x37, 0x48)),
    ], start=1):
        name = f"Heading {level}"
        if name in doc.styles:
            h = doc.styles[name]
            h.font.size = Pt(size)
            h.font.color.rgb = RGBColor(*color)
            h.font.bold = True
            h.paragraph_format.space_before = Pt(18 if level == 1 else 12)


def _parse_markdown_to_docx(doc, text):
    lines = text.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            doc.add_page_break()
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if "|" in line and not in_table:
            in_table = True
            table_rows = []

        if in_table:
            if "|" in line:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if not all(set(c) <= set("-: ") for c in cells):
                    table_rows.append(cells)
                i += 1
                continue
            else:
                if table_rows:
                    _add_table(doc, table_rows)
                in_table = False
                table_rows = []

        if line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            _add_formatted_text(p, line.strip())

        i += 1

    if table_rows:
        _add_table(doc, table_rows)


def _add_table(doc, rows):
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[i].cells[j].text = cell_text
                if i == 0:
                    for p in table.rows[i].cells[j].paragraphs:
                        for run in p.runs:
                            run.bold = True


def _add_formatted_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
