import pytest

from app.agents.ingestion import extract_text, BRIEF_SCHEMA, validate_brief


class TestExtractText:
    def test_plain_text(self):
        content = b"This is a plain text RFP document."
        result = extract_text(content, "rfp.txt")
        assert "plain text RFP" in result

    def test_pdf_extraction(self, tmp_path):
        """Create a minimal PDF and extract text."""
        import fitz
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "RFP for IT Services\nAgency: Ohio DAS\nDeadline: 2026-05-15")
        doc.save(str(pdf_path))
        doc.close()

        content = pdf_path.read_bytes()
        result = extract_text(content, "test.pdf")
        assert "RFP for IT Services" in result
        assert "Ohio DAS" in result

    def test_docx_extraction(self, tmp_path):
        """Create a minimal DOCX and extract text."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("RFP for Cloud Migration")
        doc.add_paragraph("State of California")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        content = docx_path.read_bytes()
        result = extract_text(content, "test.docx")
        assert "Cloud Migration" in result
        assert "California" in result

    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            extract_text(b"data", "file.xlsx")

    def test_empty_content_raises(self):
        with pytest.raises(ValueError, match="empty"):
            extract_text(b"", "test.pdf")


class TestBriefSchema:
    def test_schema_has_required_fields(self):
        assert "title" in BRIEF_SCHEMA
        assert "agency" in BRIEF_SCHEMA
        assert "requirements" in BRIEF_SCHEMA


class TestValidateBrief:
    def test_valid_brief(self):
        brief = {
            "title": "IT Modernization",
            "agency": "Ohio DAS",
            "state": "Ohio",
            "category": "IT",
            "deadline": "2026-05-15",
            "estimated_value": 2500000,
            "requirements": ["5 years experience"],
            "scope": "Cloud migration",
            "evaluation_criteria": {"technical": 40, "cost": 30},
        }
        result = validate_brief(brief)
        assert result["title"] == "IT Modernization"

    def test_missing_title(self):
        brief = {"agency": "Ohio DAS"}
        with pytest.raises(ValueError, match="title"):
            validate_brief(brief)

    def test_missing_agency(self):
        brief = {"title": "Test RFP"}
        with pytest.raises(ValueError, match="agency"):
            validate_brief(brief)

    def test_sets_defaults(self):
        brief = {"title": "Test", "agency": "Test Agency"}
        result = validate_brief(brief)
        assert result["requirements"] == []
        assert result["estimated_value"] is None
        assert result["scope"] == ""
        assert result["evaluation_criteria"] == {}
