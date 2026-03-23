import pytest
from app.assembly.docx_renderer import render_docx


class TestRenderDocx:
    def test_produces_docx_bytes(self):
        result = render_docx("# Test\n\nContent", title="Test")
        # DOCX is a ZIP file, starts with PK signature
        assert result[:2] == b"PK"
        assert len(result) > 100

    def test_can_reload_as_document(self):
        import io
        from docx import Document
        result = render_docx("# Heading\n\n## Sub\n\nParagraph text")
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert "Heading" in texts
        assert "Paragraph text" in texts

    def test_tables_rendered(self):
        import io
        from docx import Document
        md = "# Cost\n\n| Role | Rate |\n|------|------|\n| PM | $95 |"
        result = render_docx(md)
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) >= 1
        assert "Role" in doc.tables[0].rows[0].cells[0].text
        assert "PM" in doc.tables[0].rows[1].cells[0].text
